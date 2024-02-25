# %% [markdown]
# ### Step 0. Import all relevant libraries

# %%
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from transformers import AutoModelForQuestionAnswering, AutoTokenizer, pipeline
import nltk 
nltk.download('punkt')
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk import sent_tokenize
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.metrics.pairwise import cosine_similarity
stop_words = set(stopwords.words('english'))
from sentence_transformers import CrossEncoder
import re
import ast
import time
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path

# %% [markdown]
# ### User Inputs

# %%
TOP_N_VAL = 5

# %% [markdown]
# ### Step 1. Text Preprocessing and Partitioning

# %%
'''
Preprocessing step for splitting the legal text into sentences.
'''
alphabets= "([A-Za-z])"
prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
suffixes = "(Inc|Ltd|Jr|Sr|Co)"
starters = "(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
websites = "[.](com|net|org|io|gov)"

def split_into_sentences(text):
    text = " " + text + "  "
    text = text.replace("\n"," ")
    text = re.sub(prefixes,"\\1<prd>",text)
    text = re.sub(websites,"<prd>\\1",text)
    if "Ph.D" in text: text = text.replace("Ph.D.","Ph<prd>D<prd>")
    text = re.sub("\s" + alphabets + "[.] "," \\1<prd> ",text)
    text = re.sub(acronyms+" "+starters,"\\1<stop> \\2",text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]" + alphabets + "[.]","\\1<prd>\\2<prd>\\3<prd>",text)
    text = re.sub(alphabets + "[.]" + alphabets + "[.]","\\1<prd>\\2<prd>",text)
    text = re.sub(" "+suffixes+"[.] "+starters," \\1<stop> \\2",text)
    text = re.sub(" "+suffixes+"[.]"," \\1<prd>",text)
    text = re.sub(" " + alphabets + "[.]"," \\1<prd>",text)
    if "”" in text: text = text.replace(".”","”.")
    if "\"" in text: text = text.replace(".\"","\".")
    if "!" in text: text = text.replace("!\"","\"!")
    if "?" in text: text = text.replace("?\"","\"?")
    text = text.replace(".",".<stop>")
    text = text.replace("?","?<stop>")
    text = text.replace("!","!<stop>")
    text = text.replace("<prd>",".")
    sentences = text.split("<stop>")
    if len(sentences) > 1:
        sentences = sentences[:-1]
    sentences = [s.strip() for s in sentences]
    return sentences

def Merge(dict1, dict2):
    res = {**dict1, **dict2}
    return res

# %% [markdown]
# <h3> Step 2:- Document is pdf </h3>
# 

# %%
!pip install --upgrade langchain openai  -q
!pip install unstructured -q
!pip install unstructured[local-inference] -q
# !pip install detectron2@git+https://github.com/facebookresearch/detectron2.git@v0.6#egg=detectron2 -q
# !apt-get install poppler-utils  
from langchain.document_loaders import DirectoryLoader
directory = './content/data'

def load_docs(directory):
  loader = DirectoryLoader(directory)
  documents = loader.load()
  return documents
documents = load_docs(directory)
len(documents)
from langchain.text_splitter import RecursiveCharacterTextSplitter
law_name=" "

def split_docs(documents,chunk_size=1000,chunk_overlap=20):
  text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
  docs = text_splitter.split_documents(documents)
  return docs

docs = split_docs(documents)
print(len(docs))
print(type(docs))
var_dic ={}
count = 1
while (count <= len(docs)):
    list =[docs[count-1].page_content]
    var_dic[count] = list
    count = count + 1
print((var_dic[1]))
question = 'can an act done in good faith but without consent be an offence?'
### Step 2. Relevant Context Span Retrieval (using BCE)

#### Step 2.1. Similarity Calculation using BCE
'''
Cross Encoder - Step B
'''
def crossEncoderTopN(var_dic, QID, data_crossEncoder):
    ## Load our cross-encoder. Use fast tokenizer to speed up the tokenization
    model = CrossEncoder('cross-encoder/stsb-TinyBERT-L-4')
    
    row = data_crossEncoder.loc[data_crossEncoder["QID"] == QID]
    row = row.head(1)
    question = row.iloc[0]["Question"]
    context_spans = {}
    num_sents = 0
    
    ##Iterate over all context spans at article level
    for k in var_dic.keys():     
        passages = []
        window_size = 1        
        for l in range(len(var_dic[k])): ##Individual sub-articles (context-spans in articles)
            context = split_into_sentences(var_dic[k][l]) ##Split context span in sentences
            num_sents = num_sents + len(context)            
            for start_idx in range(0, len(context), window_size):
                end_idx = min(start_idx+window_size, len(context))
                passages.append(" ".join(context[start_idx:end_idx]))    ##Build passage of interest.
            #Concatenate the query and all passages and predict the scores for the pairs [query, passage]
            model_inputs = [[question, span] for span in passages]
            scores = model.predict(model_inputs)
            #Sort the scores in decreasing order
            results = [{'input': inp, 'score': score} for inp, score in zip(model_inputs, scores)]
            results = sorted(results, key=lambda x: x['score'], reverse=True)
            key = law_name + "_" + str(k) + "_" + str(l)
            context_spans[key] = var_dic[k][l]                        
            new_row = pd.DataFrame({
                "QID": QID,
                "Document": law_name,
                "Question": question,
                "Context_span": key,
                "CS_CrossEncoderScore": results[0]['score']
            }, index=[0])
            data_crossEncoder = pd.concat([data_crossEncoder, new_row], ignore_index=True)
    return data_crossEncoder, context_spans
### Execute Step 2.1
column_names = ['QID', 'Document', 'Question', 'Context_span', 'CS_CrossEncoderScore']
data_crossEncoder = pd.DataFrame(columns=column_names)
print(type(data_crossEncoder))
data_crossEncoder["Context_span"] = ""
data_crossEncoder["CS_CrossEncoderScore"] = 0.0
relevant_Questions = ['Q1']
df_row = {'QID': 'Q1', 'Document': law_name, 'Question': question, 'Context_span': '', 'CS_CrossEncoderScore': 0.0}
print(type(df_row))
# data_crossEncoder = data_crossEncoder.append(df_row, ignore_index = True)#

data_crossEncoder.loc[len(data_crossEncoder)] = df_row

data_crossEncoder, context_spans = crossEncoderTopN(var_dic,'Q1', data_crossEncoder)
#### Step 2.2. Ranking the context spans based on the similarity calculation in Step 2.1
'''
Ranking Process
'''
max_CrossEncoder_Scores = data_crossEncoder.groupby(['QID'], sort=False)['CS_CrossEncoderScore'].max()
data_crossEncoder['Relative_CrossEncoder'] = data_crossEncoder.apply(lambda row: row.CS_CrossEncoderScore/max_CrossEncoder_Scores[row.QID], axis = 1)
data_crossEncoder = data_crossEncoder.sort_values(["QID", "Relative_CrossEncoder"], ascending = (True, False))
data_crossEncoder["CrossEncoder_Rank"] = data_crossEncoder.groupby(['QID'], sort=False).rank(method='min',ascending=False)["Relative_CrossEncoder"]
#### Step 2.3. Top-N Context Span Selection
df_QuestionAnswering = data_crossEncoder[(data_crossEncoder.CrossEncoder_Rank <= TOP_N_VAL) & (data_crossEncoder.QID.isin(relevant_Questions))]
df_QuestionAnswering["PredictedAnswer"] = ''
df_QuestionAnswering["PredictedConfidence"] = 0.0
df_QuestionAnswering
### Step 3. Answer Extraction (using RoBERTa)
model = 'Roberta'
model_name = "deepset/roberta-base-squad2-distilled"

nlp = pipeline('question-answering', model=model_name, tokenizer=model_name)
df_QuestionAnswering = df_QuestionAnswering.reset_index()

for i,r in df_QuestionAnswering.iterrows():
    question = r['Question']
    span_id = r['Context_span']
    span = context_spans[span_id]
    QA_input = {
        'question': question,
        'context': span
    }
    print(span)
    print("\n")
    res = nlp(QA_input)
    df_QuestionAnswering.at[i,'PredictedAnswer'] = res['answer']
    df_QuestionAnswering.at[i,'PredictedConfidence'] = res['score']
    df_QuestionAnswering.at[i,'PredictedAnswer_Start'] = res['start']
    df_QuestionAnswering.at[i,'PredictedAnswer_End'] = res['end']    
### Step 4. Export the Relevant Context Spans and Highlighted Answers as a Word Document 

# %%
import datetime
# Get the current time
current_time = datetime.datetime.now()
# Format the current time as a string
time_string = current_time.strftime("%Y-%m-%d_%H-%M-%S")
df_QuestionAnswering = df_QuestionAnswering.sort_values(["PredictedConfidence", "Relative_CrossEncoder"], ascending = (False, False))
for val, cnt in df_QuestionAnswering.QID.value_counts().items():
    df_subset = df_QuestionAnswering[df_QuestionAnswering['QID'] == val]
    document = Document()
    question = df_subset['Question'].iloc[0]
    document.add_heading('Ques: ' + question, 0)
    counter = 1
    for i,r in df_subset.iterrows():
        span_id = r['Context_span']
        span = context_spans[span_id]        
        predicted_answer = df_QuestionAnswering.at[i,'PredictedAnswer']
        predicted_conf = str(df_QuestionAnswering.at[i,'PredictedConfidence'])
        crossEncoder_Rank = str(df_QuestionAnswering.at[i,'CrossEncoder_Rank'])
        
        start = int(df_QuestionAnswering.at[i,'PredictedAnswer_Start'])
        end = int(df_QuestionAnswering.at[i,'PredictedAnswer_End'])
        

        # Create a new paragraph with predicted answer highlighted
        p2 = document.add_paragraph()
        p2.add_run("\n Context Span: " + str(counter) + " \n")
        counter+=1 
                
        p2.add_run(span[0:start])
        font = p2.add_run(predicted_answer).font
        font.highlight_color = WD_COLOR_INDEX.YELLOW        
        p2.add_run(span[end:])                
        p2.add_run( " \n" + "Prediction Confidence Score: " + predicted_conf + " \n")
        
    document.save(val+f'_demofibhgfbna{time_string}.docx')
    print("Document Saved - " + val+'_demo.docx')        
        

# %%



